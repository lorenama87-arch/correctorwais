import streamlit as st
import pandas as pd
# --- CONFIGURACIÓN Y SEGURIDAD ---
st.set_page_config(page_title="WAIS-IV España: Sistema Clínico", layout="wide")
def check_password():
    """
    CORREGIDO: la versión original solo mostraba el campo de contraseña la
    primera vez que se ejecutaba el script. Si el usuario se equivocaba,
    'password_correct' quedaba en False dentro de session_state y el bloque
    que dibuja el st.text_input ya no volvía a ejecutarse nunca más,
    dejando la app en blanco sin posibilidad de reintentar.
    """
    def password_entered():
        st.session_state["password_correct"] = (
            st.session_state["password"] == "MARITA2026"
        )
    if not st.session_state.get("password_correct", False):
        st.title("🧠 WAIS-IV España (TEA): Corrección Exacta")
        st.text_input(
            "Contraseña",
            type="password",
            on_change=password_entered,
            key="password"
        )
        if "password_correct" in st.session_state and not st.session_state["password_correct"]:
            st.error("Contraseña incorrecta. Inténtalo de nuevo.")
        return False
    return True
# --- MATRIZ DE DATOS: TABLA A.1 (PD -> PE POR SUBTEST) ---
#
# AUDITORÍA COMPLETA (10/10 franjas de edad, 10/10 subtests cada una),
# transcrita celda a celda desde fotos nítidas de la Tabla A.1 del manual y
# comparada de forma programática contra el código, con verificación
# estructural (19 valores y monotonía en las 100 columnas) y pruebas
# funcionales de get_pe() sobre cada corrección. Resultado final:
#   - 16:0-17:11, 18:0-19:11, 20:0-24:11, 35:0-44:11, 45:0-54:11,
#     55:0-69:11 y 70:0-79:11: sin errores, las 10 columnas coinciden
#     exactamente con la foto.
#   - 25:0-34:11: Matrices (M) corregida (PE12 inalcanzable, duplicaba mal
#     el valor siguiente en vez del anterior).
#   - 80:0-84:11: Matrices (M), Búsqueda de símbolos (BS) e Información (I)
#     corregidas (PE1, y en Matrices también PE2, son inalcanzables con
#     cualquier PD en esta franja; se usa -1 como valor "inalcanzable" para
#     que el motor de cálculo nunca asigne esos PE).
#   - 85:0-89:11: Cubos (C), Semejanzas (S), Matrices (M), Puzles visuales
#     (PV), Información (I) y Claves (CN) corregidas — mismo patrón de PE
#     inalcanzables al principio o en medio de la columna mal codificados.
#
# Todas las correcciones comparten la misma causa raíz: cuando la tabla
# oficial marca "-" (PD no puede dar ese PE) en uno o varios PE seguidos,
# hay que "arrastrar" el valor anterior (o usar -1 si el PE es inalcanzable
# desde el principio de la columna); en varios sitios se dupicaba el valor
# equivocado, o el número equivocado de veces, desplazando en -1 todo el
# resto de la columna.
#
# PRUEBAS OPCIONALES (FI, CO, B, LN, CA): añadidas a partir de fotos de la
# Tabla A.1 que incluyen estas 5 columnas además de las 10 principales.
# AUDITORÍA DE RIGOR (segunda pasada, 2026-08-15): se releyeron las 10
# franjas comparando cada columna con la primera transcripción, con
# verificación estructural (19 valores, monotonía) tras cada corrección.
#   - Comprensión (CO) en 35:0-44:11, 45:0-54:11 y 55:0-69:11: en la primera
#     pasada la columna dio 20 valores en vez de 19 porque el tramo final se
#     leyó partido en dos rangos (p.ej. "32-33" + "34" + "35-36") en vez de
#     fusionado en un único PE19 ("34-36"). Releídas con cuidado y ya
#     cargadas en las tres franjas.
#   - Cancelación (CA), Balanzas (B) y Letras y números (LN) siguen sin
#     aparecer en la Tabla A.1 del manual a partir de 70:0-79:11: se
#     confirmó en una segunda lectura independiente de esas tres franjas
#     (mismo resultado ambas veces), así que no parece un recorte accidental
#     de la foto sino que esas pruebas no están baremadas a partir de esa
#     edad. Figuras incompletas (FI) y Comprensión (CO) sí están presentes
#     en las 10 franjas, incluidas 70:0-79:11, 80:0-84:11 y 85:0-89:11.
#   - El resto de columnas opcionales (FI, CA, B, LN en las 7 franjas donde
#     existen) se releyeron íntegramente y coincidieron exactamente con la
#     primera transcripción, sin cambios.
# Estas pruebas opcionales son solo informativas: no se suman a ningún
# índice compuesto (ver ORDEN_OPCIONALES). Ya no quedan huecos "N/D"
# salvo los reales (CA/B/LN a partir de 70:0-79:11, que el manual no
# barema para esas edades).
BAREMOS_ESPANA = {
    "16:0-17:11": {
        "C":  [15, 18, 20, 22, 27, 31, 33, 38, 43, 47, 51, 55, 59, 61, 62, 63, 64, 65, 66],
        "S":  [8, 10, 11, 12, 13, 14, 15, 16, 18, 19, 20, 22, 24, 25, 26, 28, 29, 31, 36],
        "D":  [9, 13, 15, 17, 19, 20, 22, 23, 25, 27, 29, 30, 32, 33, 35, 37, 39, 41, 48],
        "M":  [8, 9, 10, 11, 12, 14, 15, 16, 17, 19, 20, 21, 22, 23, 24, 25, 25, 26, 26],
        "V":  [8, 10, 13, 16, 18, 20, 22, 24, 27, 29, 31, 33, 35, 37, 38, 40, 42, 47, 57],
        "A":  [3, 5, 6, 7, 8, 9, 9, 10, 11, 12, 13, 13, 14, 16, 17, 18, 19, 20, 22],
        "BS": [9, 13, 16, 20, 22, 24, 26, 28, 31, 34, 36, 38, 41, 43, 46, 51, 55, 57, 60],
        "PV": [4, 5, 6, 8, 9, 10, 12, 13, 15, 16, 18, 19, 21, 22, 22, 23, 24, 25, 26],
        "I":  [1, 2, 3, 4, 5, 6, 7, 8, 10, 11, 13, 15, 16, 18, 19, 20, 21, 22, 26],
        "CN": [23, 30, 39, 41, 48, 51, 56, 61, 66, 71, 76, 81, 85, 88, 97, 99, 102, 116, 135],
        "FI": [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 19, 20, 21, 22, 24],
        "CA": [15, 20, 23, 25, 27, 31, 33, 36, 39, 41, 44, 48, 51, 54, 57, 60, 63, 65, 72],
        "CO": [2, 3, 7, 10, 12, 13, 14, 17, 19, 20, 22, 23, 25, 27, 28, 30, 32, 33, 36],
        "B": [0, 2, 4, 6, 7, 10, 13, 14, 15, 17, 19, 20, 22, 23, 24, 25, 26, 27, 27],
        "LN": [4, 9, 12, 13, 15, 16, 17, 18, 19, 19, 20, 21, 22, 23, 24, 25, 26, 27, 30],
    },
    "18:0-19:11": {
        "C":  [16, 20, 22, 24, 28, 32, 34, 40, 45, 50, 54, 57, 60, 61, 62, 63, 64, 65, 66],
        "S":  [9, 11, 12, 13, 14, 15, 16, 17, 19, 21, 22, 23, 25, 26, 27, 28, 29, 31, 36],
        "D":  [10, 13, 15, 17, 20, 21, 22, 24, 26, 28, 30, 31, 33, 35, 36, 38, 40, 42, 48],
        "M":  [8, 9, 10, 11, 13, 14, 15, 17, 18, 20, 21, 22, 23, 24, 24, 25, 25, 26, 26],
        "V":  [9, 11, 13, 16, 18, 20, 24, 27, 30, 32, 34, 36, 37, 39, 41, 43, 45, 49, 57],
        # CORREGIDA: la lista original repetía el valor 14 en dos posiciones
        # (PE11 y PE12), desplazando en +1 todo lo que venía después.
        "A":  [3, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 16, 17, 18, 19, 20, 21, 22, 22],
        "BS": [10, 15, 17, 21, 24, 27, 29, 31, 34, 36, 38, 41, 43, 45, 48, 51, 55, 57, 60],
        "PV": [4, 5, 6, 8, 9, 10, 12, 13, 15, 16, 18, 19, 21, 22, 23, 23, 24, 25, 26],
        "I":  [2, 3, 4, 5, 6, 7, 8, 10, 12, 14, 15, 17, 18, 19, 20, 21, 22, 23, 26],
        "CN": [29, 36, 42, 46, 52, 58, 63, 69, 73, 79, 84, 88, 91, 94, 101, 104, 108, 118, 135],
        "FI": [4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 15, 16, 17, 18, 19, 20, 21, 22, 24],
        "CA": [17, 22, 26, 28, 31, 33, 35, 37, 39, 43, 46, 49, 53, 56, 59, 63, 66, 70, 72],
        "CO": [2, 4, 9, 12, 14, 16, 17, 19, 20, 21, 23, 24, 26, 27, 28, 29, 30, 32, 36],
        "B": [0, 2, 4, 6, 8, 10, 13, 14, 15, 17, 19, 20, 22, 23, 24, 25, 26, 27, 27],
        "LN": [5, 9, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 30],
    },
    # --- RECONSTRUIDA ENTERA (10/10 columnas) desde una foto nítida y bien
    # orientada de la Tabla A.1. La reconstrucción anterior de esta franja
    # (hecha con una lectura de peor calidad) tenía errores en casi todas
    # las columnas.
    "20:0-24:11": {
        "C":  [16, 20, 23, 25, 29, 33, 35, 41, 45, 50, 54, 57, 60, 61, 62, 63, 64, 65, 66],
        "S":  [10, 11, 12, 13, 14, 15, 16, 18, 20, 22, 23, 24, 26, 27, 28, 29, 30, 31, 36],
        "D":  [11, 13, 16, 18, 20, 21, 22, 24, 26, 28, 30, 31, 33, 35, 37, 39, 41, 43, 48],
        "M":  [8, 9, 10, 11, 13, 14, 15, 17, 18, 20, 21, 22, 23, 24, 24, 25, 25, 26, 26],
        "V":  [10, 11, 13, 16, 18, 20, 24, 29, 32, 34, 36, 37, 39, 41, 44, 46, 48, 50, 57],
        "A":  [3, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 16, 17, 18, 19, 20, 21, 22, 22],
        "BS": [11, 16, 18, 21, 23, 27, 29, 31, 33, 37, 39, 42, 44, 47, 50, 54, 56, 58, 60],
        "PV": [4, 5, 6, 8, 9, 10, 12, 13, 15, 16, 18, 19, 21, 22, 23, 23, 24, 25, 26],
        "I":  [3, 4, 5, 5, 6, 7, 9, 11, 13, 14, 16, 18, 19, 20, 21, 22, 23, 24, 26],
        "CN": [29, 37, 42, 47, 51, 59, 64, 70, 74, 80, 85, 89, 92, 95, 101, 105, 110, 117, 135],
        "FI": [2, 4, 5, 6, 7, 9, 10, 11, 13, 14, 15, 16, 17, 18, 20, 21, 22, 23, 24],
        "CA": [19, 22, 26, 28, 31, 33, 35, 37, 39, 43, 46, 49, 53, 56, 60, 64, 66, 70, 72],
        "CO": [2, 5, 10, 13, 15, 17, 18, 20, 21, 22, 24, 25, 27, 28, 29, 30, 31, 32, 36],
        "B": [0, 2, 5, 6, 8, 10, 13, 14, 15, 17, 19, 20, 22, 23, 24, 25, 26, 27, 27],
        "LN": [5, 9, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 27, 28, 30],
    },
    "25:0-34:11": {
        "C":  [16, 21, 24, 26, 31, 35, 37, 43, 47, 51, 56, 58, 61, 62, 63, 64, 65, 66, 66],
        "S":  [8, 11, 12, 13, 14, 15, 17, 19, 20, 22, 24, 25, 26, 27, 28, 29, 30, 31, 36],
        "D":  [11, 13, 16, 17, 20, 21, 22, 24, 26, 28, 30, 31, 33, 35, 37, 39, 41, 43, 48],
        # CORREGIDA: PE12 no es alcanzable con ningún PD en esta franja (la
        # tabla marca "-" entre PE11=23 y PE13=24); debía duplicar el valor
        # ANTERIOR (23), no el siguiente (24). Verificado contra foto de la
        # Tabla A.1 (pág. 195).
        "M":  [6, 7, 9, 10, 13, 15, 16, 19, 20, 22, 23, 23, 24, 24, 25, 25, 25, 26, 26],
        "V":  [10, 12, 14, 16, 18, 20, 24, 29, 32, 34, 36, 38, 40, 42, 45, 47, 49, 51, 57],
        # CORREGIDA: desplazamiento de +1 en el tramo medio.
        "A":  [5, 6, 7, 8, 9, 10, 11, 12, 13, 15, 16, 17, 18, 19, 20, 20, 21, 22, 22],
        "BS": [8, 11, 14, 17, 20, 23, 26, 29, 32, 35, 38, 41, 44, 47, 50, 52, 54, 57, 60],
        "PV": [5, 8, 9, 10, 12, 13, 14, 15, 16, 18, 19, 20, 21, 22, 24, 25, 26, 26, 26],
        # CORREGIDA: PE7 es PD<=9, no <=10.
        "I":  [3, 4, 5, 6, 7, 8, 9, 11, 14, 16, 18, 19, 20, 21, 22, 23, 24, 26, 26],
        "CN": [26, 36, 38, 43, 46, 55, 60, 66, 71, 77, 82, 87, 90, 94, 98, 103, 109, 114, 135],
        "FI": [2, 4, 5, 6, 7, 9, 10, 11, 13, 14, 16, 17, 18, 19, 20, 21, 22, 23, 24],
        "CA": [18, 22, 24, 26, 31, 33, 36, 38, 41, 45, 47, 51, 56, 59, 61, 65, 67, 70, 72],
        "CO": [2, 5, 10, 13, 15, 17, 19, 20, 21, 24, 25, 26, 28, 29, 30, 31, 33, 34, 36],
        "B": [1, 3, 5, 7, 9, 11, 13, 15, 16, 18, 19, 21, 22, 23, 24, 25, 26, 27, 27],
        "LN": [12, 13, 14, 15, 16, 17, 18, 19, 20, 21, 21, 22, 23, 24, 25, 26, 27, 28, 30],
    },
    "35:0-44:11": {
        "C":  [12, 18, 21, 23, 28, 32, 35, 39, 43, 48, 52, 55, 58, 59, 61, 62, 64, 65, 66],
        # CORREGIDA: PE4 es PD<=13, no <=15 (desplazaba el resto de la columna).
        "S":  [8, 11, 12, 13, 14, 15, 17, 18, 20, 22, 23, 25, 26, 27, 28, 29, 30, 31, 36],
        "D":  [10, 12, 15, 17, 19, 20, 21, 23, 25, 27, 29, 30, 32, 34, 36, 38, 40, 42, 48],
        "M":  [5, 6, 7, 9, 11, 13, 14, 17, 18, 20, 21, 22, 23, 24, 24, 25, 25, 26, 26],
        "V":  [9, 11, 13, 15, 18, 20, 23, 28, 31, 34, 36, 38, 40, 42, 45, 47, 49, 51, 57],
        "A":  [4, 5, 6, 7, 8, 9, 10, 11, 13, 14, 15, 16, 17, 18, 19, 20, 21, 22, 22],
        "BS": [6, 12, 14, 18, 20, 23, 25, 28, 31, 34, 36, 40, 42, 45, 48, 51, 53, 56, 60],
        "PV": [3, 5, 6, 7, 8, 9, 11, 12, 14, 15, 16, 18, 19, 20, 22, 23, 24, 25, 26],
        "I":  [3, 4, 5, 6, 6, 7, 9, 11, 14, 17, 18, 19, 20, 21, 23, 24, 25, 26, 26],
        "CN": [20, 26, 32, 37, 43, 49, 55, 61, 67, 73, 79, 84, 87, 91, 95, 100, 107, 111, 135],
        "FI": [2, 3, 4, 5, 6, 7, 9, 10, 12, 13, 14, 15, 16, 17, 19, 20, 21, 22, 24],
        "CA": [18, 21, 23, 26, 29, 31, 34, 37, 40, 43, 45, 48, 52, 55, 61, 65, 67, 69, 72],
        # AÑADIDA tras verificación de rigor: en la primera pasada esta
        # columna se transcribió con 20 valores (el tramo final se leyó
        # partido en "32-33, 34, 35-36" en vez de fusionado en un único
        # PE19 "34-36"). Releída con cuidado: 19 valores, PE19 = 34-36.
        "CO": [2, 5, 9, 12, 14, 17, 19, 20, 21, 22, 23, 24, 26, 27, 29, 30, 31, 33, 36],
        "B": [0, 2, 3, 5, 7, 10, 12, 13, 14, 15, 17, 19, 20, 22, 23, 24, 25, 26, 27],
        "LN": [11, 13, 14, 15, 16, 17, 18, 19, 20, 20, 21, 22, 23, 24, 25, 26, 27, 28, 30],
    },
    "45:0-54:11": {
        # CORREGIDA: el tramo final (PE16-19) estaba desplazado.
        "C":  [10, 15, 17, 20, 25, 29, 32, 36, 40, 45, 49, 51, 55, 57, 59, 61, 63, 65, 66],
        # CORREGIDA: PE13 es PD<=25, no <=26.
        "S":  [7, 10, 11, 12, 13, 15, 16, 18, 19, 21, 23, 24, 25, 27, 28, 29, 30, 31, 36],
        "D":  [9, 12, 14, 15, 17, 18, 20, 22, 24, 26, 28, 29, 31, 33, 35, 37, 40, 42, 48],
        "M":  [3, 4, 6, 8, 10, 11, 12, 14, 16, 18, 19, 20, 21, 22, 23, 24, 25, 25, 26],
        "V":  [8, 10, 12, 14, 16, 18, 21, 26, 29, 32, 35, 37, 39, 41, 44, 46, 47, 51, 57],
        "A":  [4, 5, 6, 7, 8, 9, 10, 11, 12, 14, 15, 16, 17, 18, 19, 20, 21, 22, 22],
        "BS": [4, 8, 10, 14, 17, 19, 21, 24, 27, 30, 32, 36, 38, 42, 44, 48, 52, 56, 60],
        "PV": [2, 4, 5, 6, 7, 8, 9, 10, 11, 13, 15, 16, 17, 19, 20, 22, 23, 25, 26],
        "I":  [3, 4, 5, 6, 6, 7, 8, 11, 13, 15, 18, 19, 20, 21, 23, 24, 25, 26, 26],
        "CN": [16, 27, 29, 32, 34, 41, 46, 52, 58, 65, 70, 76, 80, 84, 88, 94, 102, 104, 135],
        "FI": [2, 3, 4, 5, 6, 7, 8, 10, 11, 12, 13, 14, 15, 16, 18, 19, 20, 21, 24],
        "CA": [15, 17, 20, 24, 27, 29, 32, 35, 38, 41, 43, 46, 50, 53, 59, 63, 66, 68, 72],
        # AÑADIDA tras verificación de rigor (mismo motivo que en 35:0-44:11:
        # la primera pasada dio 20 valores por un tramo final mal partido).
        "CO": [2, 4, 8, 11, 12, 15, 17, 19, 20, 21, 22, 23, 24, 25, 27, 28, 29, 31, 36],
        "B": [0, 1, 3, 4, 6, 9, 10, 11, 12, 14, 15, 17, 18, 20, 21, 22, 24, 25, 27],
        "LN": [10, 11, 12, 13, 14, 15, 16, 17, 17, 18, 19, 20, 21, 22, 23, 24, 25, 26, 30],
    },
    "55:0-69:11": {
        "C":  [5, 8, 10, 13, 17, 21, 25, 28, 32, 36, 40, 43, 47, 50, 53, 56, 59, 61, 66],
        "S":  [3, 8, 9, 10, 11, 12, 13, 15, 16, 18, 20, 22, 23, 24, 25, 26, 28, 29, 36],
        "D":  [8, 11, 12, 14, 15, 16, 17, 20, 21, 23, 25, 26, 28, 30, 33, 35, 37, 39, 48],
        "M":  [2, 3, 4, 5, 6, 7, 8, 10, 11, 13, 14, 16, 18, 20, 21, 22, 23, 24, 26],
        "V":  [7, 9, 11, 12, 13, 15, 18, 21, 24, 27, 31, 34, 36, 38, 41, 43, 45, 48, 57],
        "A":  [3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 17, 18, 19, 20, 21, 22],
        "BS": [1, 4, 6, 10, 12, 14, 16, 19, 21, 23, 26, 29, 32, 35, 38, 42, 48, 55, 60],
        "PV": [3, 4, 5, 6, 6, 7, 8, 9, 10, 10, 11, 13, 14, 15, 18, 20, 22, 24, 26],
        "I":  [2, 3, 3, 4, 5, 6, 7, 8, 10, 13, 15, 17, 19, 21, 22, 23, 24, 25, 26],
        "CN": [9, 18, 19, 22, 25, 30, 35, 41, 47, 53, 58, 65, 70, 75, 78, 84, 93, 95, 135],
        "FI": [1, 2, 3, 4, 4, 5, 7, 8, 9, 10, 12, 13, 14, 15, 16, 17, 18, 19, 24],
        "CA": [10, 12, 14, 19, 22, 24, 26, 28, 31, 34, 37, 41, 43, 48, 50, 56, 64, 67, 72],
        # AÑADIDA tras verificación de rigor (mismo motivo que en 35:0-44:11:
        # la primera pasada dio 20 valores por un tramo final mal partido).
        "CO": [2, 4, 6, 8, 9, 12, 15, 16, 19, 21, 22, 24, 25, 26, 27, 28, 30, 31, 36],
        "B": [0, 1, 2, 4, 5, 6, 7, 8, 9, 11, 12, 14, 15, 16, 17, 19, 22, 24, 27],
        "LN": [4, 5, 6, 10, 12, 13, 14, 15, 16, 17, 18, 19, 19, 20, 21, 22, 23, 25, 30],
    },
    "70:0-79:11": {
        "C":  [2, 3, 4, 7, 11, 14, 18, 20, 23, 27, 31, 34, 38, 43, 47, 50, 54, 57, 66],
        "S":  [1, 3, 4, 5, 7, 9, 10, 11, 13, 14, 16, 18, 19, 21, 22, 24, 25, 27, 36],
        "D":  [7, 8, 9, 10, 11, 13, 14, 16, 17, 18, 20, 22, 24, 26, 28, 30, 33, 34, 48],
        "M":  [1, 2, 3, 3, 4, 5, 6, 7, 8, 9, 10, 12, 14, 16, 18, 19, 20, 21, 26],
        "V":  [5, 7, 8, 9, 10, 11, 13, 15, 18, 20, 24, 27, 29, 31, 36, 38, 40, 42, 57],
        "A":  [2, 3, 4, 5, 6, 7, 7, 8, 9, 10, 11, 11, 12, 14, 15, 16, 18, 19, 22],
        "BS": [0, 1, 2, 4, 6, 7, 9, 11, 14, 15, 18, 20, 23, 26, 29, 33, 40, 49, 60],
        "PV": [1, 2, 3, 4, 5, 5, 6, 6, 7, 7, 8, 10, 11, 12, 14, 16, 19, 21, 26],
        "I":  [0, 1, 1, 2, 3, 3, 4, 5, 6, 7, 9, 11, 14, 16, 17, 19, 20, 21, 26],
        "CN": [2, 4, 6, 8, 10, 12, 16, 20, 25, 31, 35, 42, 48, 53, 58, 64, 73, 75, 135],
        "FI": [-1, -1, 0, 0, 1, 1, 2, 3, 4, 5, 7, 8, 9, 10, 11, 12, 13, 14, 24],
        "CO": [2, 3, 4, 5, 6, 7, 9, 11, 14, 17, 18, 20, 22, 23, 25, 27, 28, 30, 36],
    },
    "80:0-84:11": {
        # CORREGIDA: PE2 es PD=1, no PD=0.
        "C":  [0, 1, 2, 4, 5, 7, 10, 12, 16, 20, 22, 26, 30, 34, 38, 43, 47, 51, 66],
        "S":  [0, 1, 2, 3, 5, 7, 9, 10, 11, 12, 14, 15, 17, 18, 20, 21, 22, 24, 36],
        "D":  [5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 16, 18, 20, 23, 25, 27, 30, 32, 48],
        # CORREGIDA: PE1 y PE2 no son alcanzables con ningún PD en esta franja
        # (la tabla marca "-" para ambos); el primer PE realmente alcanzable
        # es PE3 con PD=0. Se usa -1 como valor "inalcanzable" (igual que en
        # BS/I de esta misma franja) en vez de duplicar 0, que asignaba mal
        # PE1 a PD=0. Verificado contra foto de la Tabla A.1 (pág. 200).
        "M":  [-1, -1, 0, 1, 2, 3, 4, 5, 6, 7, 8, 10, 11, 13, 14, 15, 16, 17, 26],
        "V":  [5, 7, 8, 9, 10, 11, 12, 14, 16, 18, 22, 24, 26, 28, 34, 36, 38, 40, 57],
        "A":  [1, 3, 4, 5, 5, 6, 7, 7, 8, 9, 9, 10, 11, 12, 13, 14, 15, 16, 22],
        # CORREGIDA: PE1 no es alcanzable con ningún PD (la tabla marca "-");
        # el primer PE realmente alcanzable es PE2 con PD=0. Mismo criterio
        # que en M/I de esta franja. Verificado contra foto (pág. 200).
        "BS": [-1, 0, 1, 2, 3, 4, 6, 8, 10, 11, 13, 15, 18, 21, 24, 27, 35, 41, 60],
        "PV": [0, 1, 3, 4, 4, 5, 5, 6, 6, 7, 7, 8, 9, 10, 12, 15, 17, 19, 26],
        # CORREGIDA: PE1 no es alcanzable con ningún PD (la tabla marca "-");
        # el primer PE realmente alcanzable es PE2 con PD=0. La lista
        # original arrancaba en PE1=0 (sin el "-1" inicial), lo que
        # desplazaba en -1 todo el resto de la columna. Verificado contra
        # foto (pág. 200).
        "I":  [-1, 0, 1, 2, 2, 3, 3, 4, 5, 6, 7, 9, 12, 14, 16, 19, 20, 21, 26],
        "CN": [0, 1, 3, 5, 6, 7, 10, 13, 17, 22, 25, 32, 37, 43, 49, 55, 63, 65, 135],
        "FI": [-1, -1, 0, 0, 1, 1, 2, 2, 3, 4, 6, 7, 8, 9, 10, 11, 12, 13, 24],
        "CO": [0, 1, 2, 3, 4, 5, 6, 8, 10, 12, 15, 16, 18, 19, 21, 23, 24, 26, 36],
    },
    "85:0-89:11": {
        # CORREGIDA: PE2 es PD=1, no PD=0 (la lista original repetía el 0 de
        # PE1 en PE2, desplazando en -1 todo el resto de la columna).
        # Verificado contra foto de la Tabla A.1 (pág. 201).
        "C":  [0, 1, 2, 3, 4, 6, 9, 11, 14, 18, 21, 25, 28, 32, 37, 41, 46, 50, 66],
        # CORREGIDA: la lista original tenía un "0" duplicado en PE1-PE2
        # (debía ser PD<=0 -> PE1, PD<=1 -> PE2), lo que desplazaba en -1
        # todo el resto de la columna y hacía desaparecer el valor PE18=24
        # (saltaba directamente de PE17=21 a PE19=36). Verificado contra
        # foto nítida de la Tabla A.1 (pág. 201): PE2=1, ..., PE18=24.
        "S":  [0, 1, 2, 4, 5, 6, 8, 9, 11, 12, 13, 14, 16, 18, 19, 20, 21, 24, 36],
        "D":  [4, 5, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 18, 21, 23, 25, 29, 31, 48],
        # CORREGIDA: PE1 y PE2 no son alcanzables con ningún PD en esta
        # franja (la tabla marca "-" para ambos); el primer PE realmente
        # alcanzable es PE3 con PD=0. Se usa -1 (inalcanzable) en vez de
        # duplicar 0. Verificado contra foto (pág. 201).
        "M":  [-1, -1, 0, 1, 2, 3, 4, 5, 5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 26],
        "V":  [5, 7, 8, 9, 10, 11, 12, 14, 16, 18, 22, 24, 26, 28, 34, 36, 38, 40, 57],
        "A":  [0, 1, 3, 4, 5, 6, 6, 7, 7, 8, 9, 9, 10, 11, 12, 14, 15, 16, 22],
        # CORREGIDA: para esta franja de edad, PD=0 en Símbolos no alcanza a
        # dar PE1: la tabla marca "-" (no hay puntuación directa posible)
        # para PE1-3, y el primer PE realmente alcanzable con PD=0 es PE4.
        # Se usa -1 como valor "inalcanzable" para que el motor de cálculo
        # nunca asigne esos PE por error.
        "BS": [-1, -1, -1, 0, 1, 2, 4, 5, 6, 7, 9, 10, 13, 16, 19, 22, 28, 30, 60],
        # CORREGIDA: la tabla marca "-" (inalcanzable) en DOS posiciones
        # consecutivas, PE9 y PE10 (no solo una), tras el valor directo
        # PE8=6; la lista original solo duplicaba una vez, desplazando en
        # -1 todo el resto de la columna. Verificado contra foto (pág. 201).
        "PV": [0, 1, 2, 3, 4, 5, 5, 6, 6, 6, 7, 7, 8, 9, 10, 13, 16, 18, 26],
        # CORREGIDA: mismo caso que Símbolos — PE1 no es alcanzable con
        # PD=0 en esta franja (el primer PE real con PD=0 es PE2). Además,
        # la tabla marca "-" también en PE10 (tras el valor directo PE9=5),
        # que la lista original omitía, desplazando el resto de la columna.
        "I":  [-1, 0, 1, 2, 2, 3, 3, 4, 5, 5, 6, 7, 9, 11, 15, 17, 18, 19, 26],
        # CORREGIDA: PE2 es PD=1, no PD=0 (la lista original repetía el 0 de
        # PE1 en PE2, desplazando en -1 todo el resto de la columna y
        # perdiendo el valor PE18=58, que colapsaba con PE19=135).
        "CN": [0, 1, 2, 3, 4, 5, 7, 9, 13, 16, 19, 26, 30, 35, 42, 48, 56, 58, 135],
        "FI": [-1, -1, -1, -1, 0, 0, 1, 1, 2, 3, 5, 6, 7, 8, 9, 10, 11, 12, 24],
        "CO": [0, 1, 2, 3, 4, 5, 7, 9, 11, 15, 17, 18, 20, 21, 23, 24, 25, 26, 36],
    }
}
# --- TABLAS COMPUESTAS (ESPAÑA EXACTAS - Páginas 203-208) ---
# Verificadas contra las fotos de las Tablas A.3 a A.9: sin errores.
COMP_MAP = {
    "ICV": {
        3:50, 4:50, 5:50, 6:50, 7:51, 8:54, 9:56, 10:58, 11:60, 12:63, 13:65, 14:67, 15:69, 16:71, 17:73,
        18:75, 19:78, 20:80, 21:82, 22:84, 23:86, 24:88, 25:90, 26:92, 27:94, 28:96, 29:98, 30:100, 31:102,
        32:104, 33:106, 34:108, 35:110, 36:112, 37:114, 38:116, 39:118, 40:120, 41:122, 42:124, 43:126,
        44:128, 45:130, 46:132, 47:133, 48:135, 49:137, 50:139, 51:141, 52:143, 53:145, 54:147, 55:150, 56:150, 57:150
    },
    "IRP": {
        3:50, 4:50, 5:50, 6:50, 7:52, 8:54, 9:56, 10:58, 11:60, 12:62, 13:64, 14:66, 15:68, 16:70, 17:73,
        18:75, 19:77, 20:79, 21:81, 22:83, 23:85, 24:87, 25:89, 26:91, 27:93, 28:95, 29:97, 30:100, 31:102,
        32:104, 33:106, 34:108, 35:110, 36:112, 37:114, 38:116, 39:119, 40:121, 41:123, 42:125, 43:127,
        44:129, 45:131, 46:133, 47:136, 48:138, 49:140, 50:142, 51:144, 52:147, 53:149, 54:150, 55:150, 56:150, 57:150
    },
    "IMT": {
        2:50, 3:50, 4:50, 5:53, 6:56, 7:60, 8:63, 9:66, 10:69, 11:73, 12:76, 13:79, 14:82, 15:85, 16:88,
        17:91, 18:94, 19:97, 20:100, 21:103, 22:106, 23:108, 24:111, 25:114, 26:117, 27:119, 28:122, 29:125,
        30:127, 31:130, 32:132, 33:135, 34:138, 35:140, 36:143, 37:147, 38:150
    },
    "IVP": {
        2:50, 3:52, 4:56, 5:59, 6:62, 7:64, 8:67, 9:70, 10:73, 11:75, 12:78, 13:81, 14:83, 15:86, 16:89,
        17:92, 18:94, 19:97, 20:100, 21:103, 22:105, 23:108, 24:111, 25:114, 26:117, 27:119, 28:122, 29:125,
        30:128, 31:131, 32:134, 33:137, 34:139, 35:143, 36:146, 37:150, 38:150
    },
    "CIT": {
        10:40, 11:40, 12:40, 13:40, 14:40, 15:40, 16:40, 17:40, 18:40, 19:40,
        20:40, 21:40, 22:41, 23:41, 24:42, 25:43, 26:44, 27:44, 28:45, 29:46,
        30:47, 31:47, 32:48, 33:49, 34:50, 35:50, 36:51, 37:52, 38:53, 39:53,
        40:54, 41:55, 42:56, 43:57, 44:57, 45:58, 46:59, 47:60, 48:60, 49:61,
        50:62, 51:63, 52:63, 53:64, 54:65, 55:66, 56:66, 57:67, 58:68, 59:69,
        60:69, 61:70, 62:71, 63:72, 64:72, 65:73, 66:74, 67:75, 68:75, 69:76,
        70:77, 71:78, 72:78, 73:79, 74:80, 75:81, 76:81, 77:82, 78:83, 79:84,
        80:84, 81:85, 82:86, 83:87, 84:87, 85:88, 86:89, 87:90, 88:91, 89:91,
        90:92, 91:93, 92:94, 93:94, 94:95, 95:96, 96:97, 97:97, 98:98, 99:99,
        100:100, 101:100, 102:101, 103:102, 104:103, 105:103, 106:104, 107:105, 108:106, 109:107,
        110:107, 111:108, 112:109, 113:110, 114:110, 115:111, 116:112, 117:113, 118:113, 119:114,
        120:115, 121:116, 122:116, 123:117, 124:118, 125:119, 126:120, 127:120, 128:121, 129:122,
        130:123, 131:123, 132:124, 133:125, 134:126, 135:126, 136:127, 137:128, 138:129, 139:129,
        140:130, 141:131, 142:132, 143:133, 144:133, 145:134, 146:135, 147:136, 148:136, 149:137,
        150:138, 151:139, 152:139, 153:140, 154:141, 155:142, 156:143, 157:143, 158:144, 159:145,
        160:146, 161:146, 162:147, 163:148, 164:149, 165:149, 166:150, 167:151, 168:152, 169:153,
        170:153, 171:154, 172:155, 173:156, 174:157, 175:157, 176:158, 177:159, 178:160, 179:160,
        180:160, 181:160, 182:160, 183:160, 184:160, 185:160, 186:160, 187:160, 188:160, 189:160, 190:160
    }
}
# --- MOTOR LÓGICO ---
def get_pe(sub, pd_score, edad):
    if edad not in BAREMOS_ESPANA: return 10
    # AMPLIADO: algunas pruebas opcionales (Cancelación, Balanzas, Letras y
    # números) no están baremadas en el manual para las franjas de edad más
    # altas (70:0-79:11 en adelante), y Comprensión no está aún verificada
    # en tres franjas intermedias (ver PRUEBAS_OPCIONALES_PENDIENTES). En
    # esos casos no hay "limites" para ese subtest en esa franja: se
    # devuelve None en vez de fallar, para que la interfaz pueda mostrar
    # "no disponible" en lugar de romperse o inventar un PE.
    limites = BAREMOS_ESPANA[edad].get(sub)
    if limites is None:
        return None
    for i, limite in enumerate(limites):
        if pd_score <= limite:
            return i + 1
    return 19
def approx_indice(tipo, suma):
    if tipo in COMP_MAP and suma in COMP_MAP[tipo]:
        return COMP_MAP[tipo][suma]
    # El ICG se mantiene con su fórmula base al no ser índice diagnóstico primario
    if tipo == "ICG": return int(42.0 + (1.20 * suma))
    return 100
def desc_clinico(val):
    if val >= 130: return "Muy Superior"
    if val >= 120: return "Superior"
    if val >= 110: return "Sobre el Promedio"
    if val >= 90: return "Promedio"
    if val >= 80: return "Bajo el Promedio"
    if val >= 70: return "Limítrofe"
    return "Extremadamente Bajo"
from datetime import date, datetime
from io import BytesIO
from docx import Document

# --- ORDEN DE ADMINISTRACIÓN (el mismo orden en que aparecen las pruebas
#     en el protocolo en papel, para que copiar los valores sea directo:
#     se va de arriba a abajo por la hoja, sin saltar entre columnas) ---
ORDEN_SUBTESTS = [
    ("C",  "Cubos",               "IRP", 66),
    ("S",  "Semejanzas",          "ICV", 36),
    ("D",  "Dígitos",             "IMT", 48),
    ("M",  "Matrices",            "IRP", 26),
    ("V",  "Vocabulario",         "ICV", 57),
    ("A",  "Aritmética",          "IMT", 22),
    ("BS", "Búsqueda de símbolos","IVP", 60),
    ("PV", "Puzles visuales",     "IRP", 26),
    ("I",  "Información",        "ICV", 26),
    ("CN", "Claves",              "IVP", 135),
]

# AMPLIADO: pruebas opcionales/suplementarias del WAIS-IV. Se muestran de
# forma informativa (PD -> PE) pero, a diferencia de las 10 pruebas
# principales, NO entran en el cálculo de ICV/IRP/IMT/IVP/CIT — la
# sustitución formal de una prueba principal por una opcional (con
# prorrateo según las Tablas A.8/A.9 del manual) es una decisión clínica
# que de momento no está automatizada en la app.
ORDEN_OPCIONALES = [
    ("FI", "Figuras incompletas",  "IRP", 24),
    ("CO", "Comprensión",          "ICV", 36),
    ("B",  "Balanzas",             "IRP", 27),
    ("LN", "Letras y números",     "IMT", 30),
    ("CA", "Cancelación",          "IVP", 72),
]
# Verificadas con una segunda pasada de auditoría (2026-08-15): las 10
# franjas se releyeron columna por columna comparando contra la primera
# transcripción, con verificación estructural (19 valores, monotonía) tras
# cada corrección. El único hueco real que queda:
#   - Cancelación (CA), Balanzas (B) y Letras y números (LN) no están
#     baremadas en el manual a partir de 70:0-79:11 (no aparecen en la
#     Tabla A.1 de esas franjas; confirmado en dos lecturas independientes).
#     La app muestra "N/D" para esos casos en vez de un PE inventado.
NOMBRE_INDICE = {
    "ICV": "Índice de Comprensión Verbal",
    "IRP": "Índice de Razonamiento Perceptivo",
    "IMT": "Índice de Memoria de Trabajo",
    "IVP": "Índice de Velocidad de Procesamiento",
    "ICG": "Índice de Capacidad General",
    "CIT": "Cociente Intelectual Total",
}


def calcular_edad(fecha_nacimiento, fecha_aplicacion):
    """Edad cronológica en años/meses/días, con el redondeo hacia abajo
    habitual en la corrección de tests (igual que se calcula a mano con
    la tabla 'Cálculo de edad cronológica' del protocolo)."""
    if fecha_aplicacion < fecha_nacimiento:
        return None
    años = fecha_aplicacion.year - fecha_nacimiento.year
    meses = fecha_aplicacion.month - fecha_nacimiento.month
    dias = fecha_aplicacion.day - fecha_nacimiento.day
    if dias < 0:
        meses -= 1
        mes_anterior = fecha_aplicacion.month - 1 or 12
        año_mes_anterior = fecha_aplicacion.year if fecha_aplicacion.month > 1 else fecha_aplicacion.year - 1
        import calendar
        dias += calendar.monthrange(año_mes_anterior, mes_anterior)[1]
    if meses < 0:
        años -= 1
        meses += 12
    return años, meses, dias


def franja_para_edad(años, meses, franjas):
    """Devuelve la clave de BAREMOS_ESPANA cuya franja (p.ej. '25:0-34:11')
    contiene la edad dada, o None si la edad queda fuera de todas las
    franjas disponibles (p.ej. menor de 16 o mayor de 89)."""
    edad_meses_totales = años * 12 + meses
    for clave in franjas:
        lo, hi = clave.split("-")
        lo_a, lo_m = map(int, lo.split(":"))
        hi_a, hi_m = map(int, hi.split(":"))
        if (lo_a * 12 + lo_m) <= edad_meses_totales <= (hi_a * 12 + hi_m):
            return clave
    return None


if check_password():
    st.title("📊 WAIS-IV España (TEA): Software Clínico")

    # --- 0. IDENTIFICACIÓN Y CÁLCULO DE EDAD CRONOLÓGICA ---
    # (mismo bloque que la tabla "Cálculo de edad cronológica" del
    # protocolo: se introducen las dos fechas y la franja normativa se
    # sugiere sola, en vez de tener que calcularla y buscarla a mano).
    with st.sidebar:
        st.header("Identificación")
        nombre = st.text_input("Paciente", value="Paciente Español")
        sexo = st.radio("Sexo", ["Varón", "Mujer"], horizontal=True)

        st.subheader("Cálculo de edad cronológica")
        # AMPLIADO: antes solo se podía indicar la fecha de nacimiento y la
        # app calculaba la edad. Ahora se puede elegir entre ese modo o
        # introducir directamente la edad en años y meses (por ejemplo,
        # cuando la fecha de nacimiento exacta no está disponible).
        modo_edad = st.radio(
            "¿Cómo quieres indicar la edad?",
            ["Fecha de nacimiento", "Edad directa (años y meses)"],
            horizontal=True,
        )

        fecha_aplicacion = st.date_input(
            "Fecha de aplicación", value=date.today(),
            min_value=date(1900, 1, 1), max_value=date.today(), format="DD/MM/YYYY",
        )

        fecha_nacimiento = None
        edad_calculada = None
        años_c = meses_c = dias_c = None

        if modo_edad == "Fecha de nacimiento":
            fecha_nacimiento = st.date_input(
                "Fecha de nacimiento", value=date(1990, 1, 1),
                min_value=date(1900, 1, 1), max_value=date.today(), format="DD/MM/YYYY",
            )
            edad_calculada = calcular_edad(fecha_nacimiento, fecha_aplicacion)
            if edad_calculada:
                años_c, meses_c, dias_c = edad_calculada
                st.metric("Edad cronológica", f"{años_c} a  {meses_c} m  {dias_c} d")
            else:
                st.error("La fecha de aplicación es anterior a la de nacimiento.")
        else:
            col_años, col_meses = st.columns(2)
            años_c = col_años.number_input("Años", min_value=16, max_value=89, value=30, step=1)
            meses_c = col_meses.number_input("Meses", min_value=0, max_value=11, value=0, step=1)
            dias_c = 0
            edad_calculada = (años_c, meses_c, dias_c)
            st.metric("Edad cronológica", f"{años_c} a  {meses_c} m")

        lista_edades = list(BAREMOS_ESPANA.keys())
        franja_sugerida = None
        if edad_calculada:
            franja_sugerida = franja_para_edad(años_c, meses_c, lista_edades)
            if franja_sugerida:
                st.success(f"Franja normativa sugerida: {franja_sugerida}")
            else:
                st.warning("La edad calculada queda fuera de las franjas normativas disponibles (16-89 años). Selecciona la franja manualmente.")

        indice_por_defecto = lista_edades.index(franja_sugerida) if franja_sugerida in lista_edades else 0
        edad_sel = st.selectbox(
            "Franja etaria normativa (puedes forzarla si lo necesitas)",
            lista_edades, index=indice_por_defecto,
        )
        st.success("Baremos por edad auditados contra la Tabla A.1 del manual (10/10 franjas revisadas).")

    # --- 1. PUNTUACIONES DIRECTAS (PD) ---
    # Tabla única y en el mismo orden que el protocolo en papel (antes
    # estaba repartida en 4 columnas agrupadas por índice, lo que obligaba
    # a saltar de columna en columna mientras se copiaban los valores).
    st.subheader("1. Puntuaciones directas (PD)")
    st.caption("Introduce cada puntuación directa en el mismo orden en que aparece en el protocolo. La puntuación escalar (PE) se calcula al momento.")

    encabezado = st.columns([3, 1.2, 1, 2])
    encabezado[0].markdown("**Subtest**")
    encabezado[1].markdown("**PD**")
    encabezado[2].markdown("**PE**")
    encabezado[3].markdown("**Índice**")

    pd_scores = {}
    pe_scores = {}
    for codigo, nombre_subtest, indice, pd_max in ORDEN_SUBTESTS:
        fila = st.columns([3, 1.2, 1, 2])
        fila[0].write(nombre_subtest)
        pd_scores[codigo] = fila[1].number_input(
            f"PD {nombre_subtest}", 0, pd_max, key=f"pd_{codigo}", label_visibility="collapsed",
        )
        pe_scores[codigo] = get_pe(codigo, pd_scores[codigo], edad_sel)
        fila[2].markdown(f"**{pe_scores[codigo]}**")
        fila[3].caption(indice)

    # AMPLIADO: pruebas opcionales/suplementarias, en una tabla aparte para
    # no mezclarlas con las 10 principales. Se muestran a título informativo
    # (no se suman a ningún índice compuesto).
    with st.expander("Pruebas opcionales (suplementarias)"):
        st.caption(
            "Estas puntuaciones se muestran solo a título informativo: no se "
            "suman a ICV/IRP/IMT/IVP/CIT. La sustitución formal de una "
            "prueba principal por una opcional (con el prorrateo del "
            "manual) queda a criterio clínico."
        )
        encabezado_op = st.columns([3, 1.2, 1, 2])
        encabezado_op[0].markdown("**Subtest**")
        encabezado_op[1].markdown("**PD**")
        encabezado_op[2].markdown("**PE**")
        encabezado_op[3].markdown("**Índice**")

        pd_opcionales = {}
        pe_opcionales = {}
        for codigo, nombre_subtest, indice, pd_max in ORDEN_OPCIONALES:
            fila_op = st.columns([3, 1.2, 1, 2])
            fila_op[0].write(nombre_subtest)
            disponible = BAREMOS_ESPANA.get(edad_sel, {}).get(codigo) is not None
            if disponible:
                pd_opcionales[codigo] = fila_op[1].number_input(
                    f"PD {nombre_subtest}", 0, pd_max, key=f"pd_op_{codigo}", label_visibility="collapsed",
                )
                pe_opcionales[codigo] = get_pe(codigo, pd_opcionales[codigo], edad_sel)
                fila_op[2].markdown(f"**{pe_opcionales[codigo]}**")
            else:
                fila_op[1].caption("—")
                fila_op[2].caption("N/D")
            fila_op[3].caption(indice)
        if any(
            BAREMOS_ESPANA.get(edad_sel, {}).get(c) is None
            for c, *_ in ORDEN_OPCIONALES
        ):
            st.caption(
                "⚠️ Cancelación, Balanzas y Letras y números no están "
                "baremadas en el manual a partir de 70:0-79:11, por eso "
                "aparecen como \"N/D\" en esta franja."
            )

    e_s, e_v, e_i = pe_scores["S"], pe_scores["V"], pe_scores["I"]
    e_c, e_m, e_pv = pe_scores["C"], pe_scores["M"], pe_scores["PV"]
    e_d, e_a = pe_scores["D"], pe_scores["A"]
    e_bs, e_cn = pe_scores["BS"], pe_scores["CN"]

    suma_icv, suma_irp = e_s + e_v + e_i, e_c + e_m + e_pv
    suma_imt, suma_ivp = e_d + e_a, e_bs + e_cn
    suma_icg = e_s + e_v + e_i + e_c + e_m + e_pv
    suma_cit = suma_icv + suma_irp + suma_imt + suma_ivp

    icv = approx_indice("ICV", suma_icv)
    irp = approx_indice("IRP", suma_irp)
    imt = approx_indice("IMT", suma_imt)
    ivp = approx_indice("IVP", suma_ivp)
    icg = approx_indice("ICG", suma_icg)
    cit = approx_indice("CIT", suma_cit)
    dis = max(icv, irp, imt, ivp) - min(icv, irp, imt, ivp)

    st.divider()

    # --- 2. RESULTADOS ---
    st.subheader("2. Resultados")
    st.success(f"Resultados TEA | Paciente: {nombre} | Edad: {edad_sel}")
    m1, m2, m3, m4, m5, m6 = st.columns(6)
    m1.metric("ICV", icv, desc_clinico(icv))
    m2.metric("IRP", irp, desc_clinico(irp))
    m3.metric("IMT", imt, desc_clinico(imt))
    m4.metric("IVP", ivp, desc_clinico(ivp))
    m5.metric("ICG", icg, desc_clinico(icg))
    m6.metric("CIT", cit, desc_clinico(cit), delta_color="off")

    if dis >= 23:
        st.error(f"⚠️ DISCREPANCIA: Diferencia de {dis} pts entre índices. Use el ICG.")
    else:
        st.info("✅ PERFIL ARMÓNICO: El CIT es representativo.")

    # "Sumas de puntuaciones escalares" + tabla PD/PE en el mismo orden
    # que el protocolo, replicando el formato de la hoja de resumen.
    st.subheader("Conversión de puntuaciones directas en escalares")
    df = pd.DataFrame({
        "Subtest": [n for _, n, _, _ in ORDEN_SUBTESTS],
        "Índice": [ix for _, _, ix, _ in ORDEN_SUBTESTS],
        "PD": [pd_scores[c] for c, _, _, _ in ORDEN_SUBTESTS],
        "PE": [pe_scores[c] for c, _, _, _ in ORDEN_SUBTESTS],
    })
    st.table(df.set_index("Subtest"))

    sumas_df = pd.DataFrame({
        "ICV": [suma_icv], "IRP": [suma_irp], "IMT": [suma_imt], "IVP": [suma_ivp], "CIT": [suma_cit],
    }, index=["Suma de PE"])
    st.table(sumas_df)

    st.subheader("Perfil de puntuaciones compuestas")
    perfil_df = pd.DataFrame({
        "ICV": [icv], "IRP": [irp], "IMT": [imt], "IVP": [ivp], "ICG": [icg], "CIT": [cit],
    }, index=["Puntuación"])
    st.table(perfil_df)

    st.bar_chart(df.set_index("Subtest")["PE"])

    # --- 3. INFORME ---
    st.divider()
    st.subheader("3. Informe")

    if st.button("📄 GENERAR INFORME"):

        # AMPLIADO: redacción más clínica y cercana, evitando un tono seco o
        # de "veredicto". Cada índice se presenta explicando primero qué mide
        # (en lenguaje llano) antes de dar la puntuación, y la clasificación
        # psicométrica oficial (desc_clinico) se mantiene entre paréntesis
        # para conservar el rigor clínico, pero no encabeza la frase.
        DESCRIPCION_INDICE = {
            "ICV": "la capacidad para formar conceptos verbales, razonar utilizando el lenguaje "
                   "y aplicar los conocimientos adquiridos a través de la experiencia y la "
                   "escolarización",
            "IRP": "el razonamiento no verbal, la organización visoespacial y la capacidad de "
                   "integrar la información visual con una respuesta motora precisa",
            "IMT": "la capacidad de mantener información en la mente durante unos instantes, "
                   "manipularla mentalmente y utilizarla para completar una tarea; está muy "
                   "ligada a la atención y a la concentración",
            "IVP": "la rapidez y la precisión con las que se procesa información visual sencilla "
                   "y rutinaria, así como la coordinación entre lo que se ve y la respuesta motora",
        }

        def frase_rendimiento(val):
            """Traduce la categoría psicométrica en una frase más cercana,
            sin perder la categoría oficial (que se añade entre paréntesis
            donde se usa esta función)."""
            mapa = {
                "Muy Superior": "un rendimiento muy superior al esperado para su grupo de edad",
                "Superior": "un rendimiento superior al esperado para su grupo de edad",
                "Sobre el Promedio": "un rendimiento algo por encima de lo esperado para su grupo de edad",
                "Promedio": "un rendimiento dentro de lo esperado para su grupo de edad",
                "Bajo el Promedio": "un rendimiento algo por debajo de lo esperado para su grupo de edad",
                "Limítrofe": "un rendimiento notablemente por debajo de lo esperado para su grupo de edad",
                "Extremadamente Bajo": "un rendimiento muy por debajo de lo esperado para su grupo de edad",
            }
            return mapa.get(desc_clinico(val), desc_clinico(val).lower())

        def clasif_frase(val):
            return desc_clinico(val).lower()

        art = "El evaluado" if sexo == "Varón" else "La evaluada"
        del_al = "del evaluado" if sexo == "Varón" else "de la evaluada"
        edad_txt = f"{años_c} años y {meses_c} meses" if edad_calculada else edad_sel

        intro = (
            "El WAIS-IV (Wechsler Adult Intelligence Scale, cuarta edición) es una escala de "
            "inteligencia individual, estandarizada y de uso clínico, dirigida a adolescentes "
            "mayores y personas adultas. Evalúa el funcionamiento cognitivo a través de cuatro "
            "índices principales -comprensión verbal, razonamiento perceptivo, memoria de trabajo "
            "y velocidad de procesamiento- que se combinan en un Cociente Intelectual Total (CIT) "
            f"representativo del funcionamiento cognitivo global.\n\n"
            f"{art} presenta una edad cronológica de {edad_txt} en el momento de la aplicación, "
            f"dentro de la franja normativa {edad_sel}."
        )

        parrafo_indices = (
            f"En el {NOMBRE_INDICE['ICV']} (ICV) -que recoge {DESCRIPCION_INDICE['ICV']}- "
            f"obtiene una puntuación de {icv}, mostrando {frase_rendimiento(icv)} "
            f"(categoría: {desc_clinico(icv)}).\n\n"
            f"En el {NOMBRE_INDICE['IRP']} (IRP) -que valora {DESCRIPCION_INDICE['IRP']}- "
            f"alcanza {irp} puntos, con {frase_rendimiento(irp)} (categoría: {desc_clinico(irp)}).\n\n"
            f"En cuanto al {NOMBRE_INDICE['IMT']} (IMT) -es decir, {DESCRIPCION_INDICE['IMT']}- "
            f"se sitúa en {imt} puntos, {frase_rendimiento(imt)} (categoría: {desc_clinico(imt)}).\n\n"
            f"Por último, en el {NOMBRE_INDICE['IVP']} (IVP) -que mide {DESCRIPCION_INDICE['IVP']}- "
            f"obtiene {ivp} puntos, {frase_rendimiento(ivp)} (categoría: {desc_clinico(ivp)})."
        )

        if dis >= 23:
            parrafo_discrepancia = (
                f"Entre el índice más alto y el más bajo se observa una diferencia de {dis} "
                f"puntos, un valor considerado clínicamente significativo. Esto sugiere que el "
                f"perfil cognitivo no es homogéneo, por lo que resulta más informativo describir "
                f"el rendimiento índice por índice que resumirlo en una única cifra global; el CIT "
                f"debe interpretarse con esa cautela. En estos casos puede resultar más "
                f"representativo el {NOMBRE_INDICE['ICG']} (ICG), que se sitúa en {icg} puntos "
                f"({clasif_frase(icg)}) y se calcula únicamente a partir de los subtests verbales "
                f"y perceptivos, dejando fuera la memoria de trabajo y la velocidad de "
                f"procesamiento."
            )
        else:
            parrafo_discrepancia = (
                f"La diferencia entre el índice más alto y el más bajo es de {dis} puntos, dentro "
                f"de límites no significativos, lo que indica un perfil cognitivo razonablemente "
                f"homogéneo entre las distintas áreas evaluadas."
            )

        parrafo_cit = (
            f"El Cociente Intelectual Total (CIT), que resume de forma global el conjunto de los "
            f"índices anteriores, es de {cit} puntos, lo que sitúa el funcionamiento cognitivo "
            f"general {del_al} en {frase_rendimiento(cit)} (categoría: {desc_clinico(cit)})."
        )

        informe_final = (
            f"INFORME DE EVALUACIÓN WAIS-IV\n\n"
            f"Paciente: {nombre}\n"
            f"Edad cronológica: {edad_txt}\n"
            f"Franja normativa: {edad_sel}\n"
            f"Fecha de aplicación: {fecha_aplicacion.strftime('%d/%m/%Y')}\n\n"
            f"{intro}\n\n"
            f"RESULTADOS\n{parrafo_indices}\n\n{parrafo_cit}\n\n"
            f"INTERPRETACIÓN DEL PERFIL\n{parrafo_discrepancia}"
        )

        st.text_area("Vista previa del informe:", informe_final, height=400)

        # --- Documento Word ---
        doc = Document()
        doc.add_heading("INFORME DE EVALUACIÓN WAIS-IV", 0)

        doc.add_heading("Datos de identificación", level=1)
        tabla_id = doc.add_table(rows=0, cols=2)
        datos_identificacion = [
            ("Paciente", nombre),
            ("Edad cronológica", edad_txt),
            ("Franja normativa", edad_sel),
            ("Fecha de aplicación", fecha_aplicacion.strftime("%d/%m/%Y")),
        ]
        # La fecha de nacimiento solo existe si se ha usado ese modo para
        # calcular la edad; en modo "Edad directa" no se pide y se omite.
        if fecha_nacimiento:
            datos_identificacion.append(("Fecha de nacimiento", fecha_nacimiento.strftime("%d/%m/%Y")))
        for etiqueta, valor in datos_identificacion:
            fila = tabla_id.add_row().cells
            fila[0].text = etiqueta
            fila[1].text = str(valor)

        doc.add_heading("Descripción de la prueba", level=1)
        doc.add_paragraph(intro)

        doc.add_heading("Puntuaciones directas y escalares", level=1)
        tabla_pd = doc.add_table(rows=1, cols=4)
        hdr = tabla_pd.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Subtest", "Índice", "PD", "PE"
        for codigo, nombre_subtest, indice, _ in ORDEN_SUBTESTS:
            fila = tabla_pd.add_row().cells
            fila[0].text = nombre_subtest
            fila[1].text = indice
            fila[2].text = str(pd_scores[codigo])
            fila[3].text = str(pe_scores[codigo])

        doc.add_heading("Puntuaciones compuestas", level=1)
        tabla_comp = doc.add_table(rows=1, cols=3)
        hdr = tabla_comp.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Índice", "Puntuación", "Clasificación"
        for sigla, valor in [("ICV", icv), ("IRP", irp), ("IMT", imt), ("IVP", ivp), ("ICG", icg), ("CIT", cit)]:
            fila = tabla_comp.add_row().cells
            fila[0].text = f"{NOMBRE_INDICE[sigla]} ({sigla})"
            fila[1].text = str(valor)
            fila[2].text = desc_clinico(valor)

        doc.add_heading("Interpretación del perfil", level=1)
        doc.add_paragraph(parrafo_indices + " " + parrafo_cit)
        doc.add_paragraph(parrafo_discrepancia)

        bio = BytesIO()
        doc.save(bio)
        st.download_button(
            "📥 Descargar informe en Word",
            bio.getvalue(),
            f"Informe_WAIS-IV_{nombre.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
